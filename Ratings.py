class AutomaticBrickMover:
    def __init__(self,path="C:\\Users\\12580\\Desktop\\a.docx",picpath="C:\\Users\\12580\\Desktop\\b.jpg",
                 ratingpath="C:\\Users\\12580\\Desktop\\评级调整.xlsx",bond_name=['']):
     self.path=path
     self.picpath=picpath
     self.ratingpath = ratingpath
     self.bond_name=bond_name
    def BrickMover(self):
     try:
      modify
     except NameError:
      def modify(string):
       English = u'\000'
       Chinese = u''
       table = {ord(o): ord(k) for o, k in zip(English, Chinese)}
       return string.translate(table)
     try:
      mirrormodify
     except NameError:
      def mirrormodify(string):
       English = u''
       Chinese = u''
       table = {ord(o): ord(k) for o, k in zip(Chinese, English)}
       return string.translate(table)
     try:
      from docx import Document
     except ModuleNotFoundError:
      print("需要先下载docx库")
     doc = Document(self.path)
     content=""
     for paragraph in doc.paragraphs:
        content=content+paragraph.text
     NewContent=modify(content)
     Numlist=["1","2","3","4","5","6","7","8","9","0"]
     for i in range(len(NewContent)):
      if NewContent[i-1] in Numlist and NewContent[i+1] in Numlist and i not in [0,len(NewContent)-1]:
        listNewContent=list(NewContent)
        listNewContent[i]=mirrormodify(listNewContent[i])
        NewContent="".join(listNewContent)
     doc._body.clear_content()
     doc.add_paragraph(NewContent)
     doc.save(self.path)

    def Pyteserrect(self):
     from PIL import Image
     import pytesseract
     tesseract_cmd=r"C:\Users\12580\Tesseract-OCR\tesseract.exe"
     image=Image.open(self.picpath)
     from PIL import ImageEnhance
     image = image.convert('L')
     enhancer = ImageEnhance.Color(image)
     enhancer = enhancer.enhance(0)
     enhancer = ImageEnhance.Brightness(enhancer)
     enhancer = enhancer.enhance(2)
     enhancer = ImageEnhance.Contrast(enhancer)
     enhancer = enhancer.enhance(8)
     enhancer = ImageEnhance.Sharpness(enhancer)
     image = enhancer.enhance(20)
     text=pytesseract.image_to_string(image,lang="chi_sim")
     text=text.replace(" ","")
     text=text.replace("\n", "")
     text=text.replace("\x0c", "")
     from docx import Document
     doc = Document(self.path)
     doc._body.clear_content()
     doc.add_paragraph(text)
     doc.save(self.path)
     self.BrickMover()

    def BaiduOcr(self):
        from aip import AipOcr
        from PIL import Image
        APP_ID = ""       #这几行需要自行注册百度OCR账号
        API_KEY = ""
        SECRET_KEY = ""
        client = AipOcr(APP_ID, API_KEY, SECRET_KEY)
        text = []
        try:
            getImage
        except NameError:
            def getImage(picpath):
                with open(picpath, "rb") as fp:
                    return fp.read()
        image = getImage(self.picpath)
        Newtext = client.basicAccurate(image)    #额度用完换成basicGeneral
        Newtext = Newtext['words_result']
        for i in range(len(Newtext)):
            text = text + list(Newtext[i]['words'])
        text = "".join(text)
        from docx import Document
        doc = Document(self.path)
        doc._body.clear_content()
        doc.add_paragraph(text)
        doc.save(self.path)
        self.BrickMover()

    def create_template_to_new_location(self):
        import pandas as pd
        from datetime import datetime, timedelta
        try:
            modified_monday
        except NameError:
            def modified_monday(weekday=1, d=datetime.now()):
                delta = weekday - d.isoweekday()
                if delta == -4 and int(d.strftime('%H')) >= 12:   #周五12点之后刷就会变成下下周一
                    delta += 14
                elif delta == -5 or delta == -7:
                    delta += 14
                else:
                    delta += 7
                newtime = d + timedelta(delta)
                return datetime.strftime(newtime, "%Y%m%d")
        workbook=pd.read_excel(self.ratingpath,sheet_name="评级调整")
        backup_frame=pd.DataFrame()
        for x in self.bond_name:
            small_frame=workbook[workbook.发债机构==x]
            if len(small_frame)==0:
                print("%s需要手动输入信息"%(x))
                continue
            small_frame=pd.DataFrame(small_frame.iloc[len(small_frame) - 1,:]).T
            small_frame.iloc[0,6] = small_frame.iloc[0,8]
            small_frame.iloc[0,18] = modified_monday()[2:]+"评级调整"
            backup_frame=backup_frame.append(small_frame)
        backup_frame.to_excel("C:\\Users\\12580\\Desktop\\评级准备.xls",index=False,sheet_name="评级调整")

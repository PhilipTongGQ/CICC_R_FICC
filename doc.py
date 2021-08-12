import pandas as pd
import numpy as np
import math
import datetime
from datetime import timedelta
from docx import Document

class doc_template:
    def __init__(self, targetpath="C:\\Users\\12580\\Desktop\\quick1.docx",   #这里的path都是双斜杠！！！注意根据自己电脑的路径修改
                 sourcepath="C:\\Users\\12580\\Desktop\\data.docx",
                 excelpath="C:\\Users\\12580\\Desktop\\模板准备.xlsx",
                 issuerpath="C:\\Users\\12580\\Desktop\\CP Issue Index -值格式.xlsx",
                 cover_page_path="C:\\Users\\12580\\Desktop\\封面表准备.xlsx",
                 historypath="C:\\Users\\12580\\Desktop\\00所有覆盖债券首页.xlsx"):
        self.targetpath = targetpath
        self.sourcepath = sourcepath
        self.excelpath = excelpath
        self.issuerpath = issuerpath
        self.cover_page_path=cover_page_path
        self.historypath=historypath
        self.name_switch={"中诚信国际信用评级有限责任公司":"中诚信","联合资信评估股份有限公司":"联合","大公国际资信评估有限公司":"大公","上海新世纪资信评估投资服务有限公司":"新世纪","东方金诚国际信用评估有限公司":"东方金诚","中证鹏元资信评估股份有限公司":"鹏元"}

    def search(self):
        gs_table = pd.read_excel(self.excelpath)
        fill="啥都没有"
        CPMTN_bond_history=pd.read_excel(self.historypath,sheet_name="短融中票")
        CPMTN_bond_history.iloc[:, 11] = CPMTN_bond_history.iloc[:, 11].fillna(fill)
        CB_bond_history=pd.read_excel(self.historypath,sheet_name="公司债企业债")
        CB_bond_history.iloc[:, 11] = CB_bond_history.iloc[:, 11].fillna(fill)
        CPMTN_anchor=pd.DataFrame([np.where(gs_table.iloc[i, 31] == CPMTN_bond_history.iloc[:, 20])[0] for i in range(len(gs_table))])    #这个dataframe里每一行是每个券在所有覆盖债券短融里对应行的位置
        CPMTN_final=pd.DataFrame(CPMTN_anchor.iloc[i, CPMTN_anchor.iloc[i,:].notna()[::-1].idxmax()] for i in range(len(CPMTN_anchor)))
        CPMTNarray=np.array(CPMTN_final.iloc[:,0].T).astype(int)     #不知道为啥lambda公式直接在上面dataframe转为int不行，有兴趣可以改改
        for i in range(len(CPMTN_final)):
            if CPMTNarray[i]>=0:
                if CPMTN_bond_history.iloc[CPMTNarray[i],11] == fill and CPMTN_bond_history.iloc[CPMTNarray[i], 20] == CPMTN_bond_history.iloc[CPMTNarray[i] - 1, 20]:  #同主体相连的两个债（两个品种）导致pandas读取第二个债日期为nan（被改成了str啥都没有），所以直接用上一行的日期
                    CPMTNarray[i]=CPMTNarray[i]-1
                CPMTN_final.loc[i,"最近报告日期"] = CPMTN_bond_history.iloc[CPMTNarray[i],11]
                CPMTN_final.loc[i, "行业"] = CPMTN_bond_history.iloc[CPMTNarray[i], 7]
                if isinstance(CPMTN_final.loc[i,"最近报告日期"], str):                             #下面处理脏数据，因为日期可能在12列也可能在13列。
                    CPMTN_final.loc[i, "最近报告日期"]= CPMTN_bond_history.iloc[CPMTNarray[i],12]
                    if isinstance(CPMTN_final.loc[i,"最近报告日期"], str):
                        CPMTN_final.loc[i, "最近报告日期"] = CPMTN_bond_history.iloc[CPMTNarray[i], 13]
                else: continue
            else:
                CPMTN_final.loc[i,"最近报告日期"]=np.nan
                CPMTN_final.loc[i, "行业"] = "顺便一起看了吧"
        CPMTN_final["最近报告日期"] = CPMTN_final["最近报告日期"].astype(object).where(CPMTN_final["最近报告日期"].notnull(), np.nan)
        CB_anchor=pd.DataFrame([np.where(gs_table.iloc[i, 31] == CB_bond_history.iloc[:, 20])[0] for i in range(len(gs_table))])    #这个dataframe里每一行是每个券在所有覆盖债券企业公司债里对应行的位置
        CB_final = pd.DataFrame(CB_anchor.iloc[i, CB_anchor.iloc[i,:].notna()[::-1].idxmax()] for i in range(len(CB_anchor)))
        CBarray=np.array(CB_final.iloc[:,0].T).astype(int)
        for i in range(len(CB_final)):
            if CBarray[i]>=0:
                if CB_bond_history.iloc[CBarray[i],11] == fill and CB_bond_history.iloc[CBarray[i], 20] == CB_bond_history.iloc[CBarray[i] - 1, 20]:  #同主体相连的两个债导致pandas读取第二个债日期为nan（被改成了str啥都没有），所以直接用上一行的日期
                    CBarray[i]=CBarray[i]-1
                CB_final.loc[i,"最近报告日期"] = CB_bond_history.iloc[CBarray[i],11]
                CB_final.loc[i, "行业"] = CB_bond_history.iloc[CBarray[i], 7]
                if isinstance(CB_final.loc[i,"最近报告日期"], str):   #不会出现nan，因为最前面导入bond_history的时候在空值里填入了str“啥都没有”
                    CB_final.loc[i, "最近报告日期"]= CB_bond_history.iloc[CBarray[i],12]
                    try:
                        math.isnan(CB_final.loc[i, "最近报告日期"])           #所有覆盖债券表sheet企业债1000行左右开始出现nan
                        CB_final.loc[i, "最近报告日期"] = CB_bond_history.iloc[CBarray[i], 13]   #先检测是不是nan，如果是的话这一行就会运行成功，否则就不回。
                    except TypeError:
                       if isinstance(CB_final.loc[i,"最近报告日期"], str):    #如果不是nan就试是不是更新覆盖过
                           CB_final.loc[i, "最近报告日期"] = CB_bond_history.iloc[CBarray[i], 13]
                       else:                                                #如果既不是nan(属于float)也不是str(有星号代表那一期周报重新覆盖过)，就是正常的datetime，所以什么也不用做
                          pass
                else: continue
            else:
                CB_final.loc[i,"最近报告日期"]=np.nan
                CB_final.loc[i, "行业"] = "顺便一起看了吧"
        CB_final["最近报告日期"] = CB_final["最近报告日期"].astype(object).where(CB_final["最近报告日期"].notnull(), np.nan)
        transfer_table = pd.DataFrame()
        if gs_table.iloc[0, 65].find("中")<0 and gs_table.iloc[0, 65].find("短")<0:  #周报是星期几？
            day=4    #CB
        else:
            day=3    #CPMTN
        try:
            Publication_date
        except NameError:
            def Publication_date(weekday=day, d=datetime.datetime.now()):
                delta = weekday - d.isoweekday()
                if delta == 0 and int(d.strftime('%H')) >= 12:  # 周三/周四12点之后刷就会变成下下周一
                    delta += 7
                elif delta < 0:
                    delta += 7
                return d + timedelta(delta)
        for i in range(len(CB_final)):  #或者CPMTN无所谓反正都是一样长度
            transfer_table.loc[i, "债券简称"] = gs_table.iloc[i,1]
            if isinstance(CB_final.iloc[i, 1], datetime.datetime) and isinstance(CPMTN_final.iloc[i,1], datetime.datetime):
                if CB_final.iloc[i, 1]>CPMTN_final.iloc[i,1]:
                    transfer_table.loc[i, "行业"] = CB_final.iloc[i, 2]
                    if Publication_date() - timedelta(days=180)>CB_final.iloc[i, 1]:
                        transfer_table.loc[i,"最后日期"] = CB_final.iloc[i, 1].strftime('%F')[2:]+"*（P）"
                    else:
                        transfer_table.loc[i, "最后日期"] = CB_final.iloc[i, 1].strftime('%F')[2:]
                else:
                    transfer_table.loc[i, "行业"] = CPMTN_final.iloc[i, 2]
                    if Publication_date() - timedelta(days=180)>CPMTN_final.iloc[i, 1]:
                        transfer_table.loc[i, "最后日期"] = CPMTN_final.iloc[i,1].strftime('%F')[2:]+"*（P）"
                    else:
                        transfer_table.loc[i, "最后日期"] = CPMTN_final.iloc[i,1].strftime('%F')[2:]
            elif isinstance(CB_final.iloc[i, 1], float) and isinstance(CPMTN_final.iloc[i,1], datetime.datetime):
                transfer_table.loc[i, "行业"] = CPMTN_final.iloc[i, 2]
                if Publication_date() - timedelta(days=180)>CPMTN_final.iloc[i, 1]:
                     transfer_table.loc[i, "最后日期"] = CPMTN_final.iloc[i, 1].strftime('%F')[2:]+"*（P）"
                else:
                     transfer_table.loc[i, "最后日期"] = CPMTN_final.iloc[i, 1].strftime('%F')[2:]
            elif isinstance(CPMTN_final.iloc[i, 1], float) and isinstance(CB_final.iloc[i,1], datetime.datetime):
                transfer_table.loc[i, "行业"] = CB_final.iloc[i, 2]
                if Publication_date() - timedelta(days=180)>CB_final.iloc[i, 1]:
                     transfer_table.loc[i, "最后日期"] = CB_final.iloc[i, 1].strftime('%F')[2:]+"*（P）"
                else:
                    transfer_table.loc[i, "最后日期"] = CB_final.iloc[i, 1].strftime('%F')[2:]
            else:
                transfer_table.loc[i, "行业"] = "顺便一起看了吧"
                transfer_table.loc[i, "最后日期"]="没找到，请手动查询"
        transfer_table.to_excel(self.cover_page_path, index=False)
        print("Search is over")    #搜完了！


        #为了防止返回时间因为漏券错误，同时返回最后一次出现的对应日期
        #最后查找一遍如果主体名称没有一起报错

    def get_bond_name(bondnamelist,ind,row,gs_table):
        bondnamelist += [row[x].text for x in range(len(row)) if x != 0]
        if len(bondnamelist) > 0:
            global loc
            loc = np.where(gs_table.债券简称 == bondnamelist[0])[0][0]
        else:
            print("第%d个table名字没抓到" % (ind - 1))          #这行好像没用

    def get_length(loc,row,gs_table):
        if str(gs_table.iloc[loc, 8]) is not None:    #期限应该只会是str
            if str(gs_table.iloc[loc, 8]).find("+")!=-1:
                if str(gs_table.iloc[loc, 8]).find("N")!=-1:
                    row[1].text = gs_table.iloc[loc, 8]
                elif len(str(gs_table.iloc[loc, 8]))>=5 and str(gs_table.iloc[loc, 8]).find("+")!=-1:
                    row[1].text = str(gs_table.iloc[loc, 7])+"Y"+str(gs_table.iloc[loc, 8])[0]+"P"+ str(int(gs_table.iloc[loc, 8][0])+int(gs_table.iloc[loc, 8][2]))+"P"
                else:
                    row[1].text=str(gs_table.iloc[loc, 7])+"Y"+str(gs_table.iloc[loc, 8][0])+"P"
            elif str(gs_table.iloc[loc, 8]).find("D")!=-1:
                row[1].text = str(gs_table.iloc[loc, 7])+"年"
            else: row[1].text = str(gs_table.iloc[loc, 7])+"年"
        else:
            row[1].text = str(gs_table.iloc[loc, 7])+"年"

    def get_credit_boost(self,row,loc,gs_table):   #增信措施，起个炫酷的英文名吧
        try:
            math.isnan(gs_table.iloc[loc, 27])
            row[1].text = ""
        except TypeError:
            print("%s有担保，请检查是否为抵质押担保"%(gs_table.iloc[loc,1]))
            row[1].text = str(gs_table.iloc[loc,37])

    def get_ratings(self,row,gs_table,type):
        issuer_rating = gs_table.iloc[loc, 10]
        rating_agency = str(gs_table.iloc[loc, 28])
        for key, items in self.name_switch.items():
            rating_agency = rating_agency.replace(key, items)
        if type=="CB":
            bond_rating = gs_table.iloc[loc, 9]
            row[1].text = rating_agency + "，" + issuer_rating + "/" + bond_rating
        if type=="CPMTN":
            try:
                math.isnan(gs_table.iloc[loc, 9])
                row[1].text = rating_agency + "评定为NA" + "，" + "主体评级" + issuer_rating
            except TypeError:
                bond_rating = str(gs_table.iloc[loc, 9])
                row[1].text = rating_agency + "评定为%s" % (bond_rating) + "，" + "主体评级" + issuer_rating

    def get_CICC_ratings(row,loc,gs_table,issuer_column,issuer):
        if len(np.where(issuer_column == doc_template.parentheses(gs_table.iloc[loc, 31]))[0]) == 0:
            row[1].text = "手动查找"
            print("%s的主体全称在issue表未找到" % (gs_table.iloc[loc, 1]))
        else:
            location = np.where(issuer_column == doc_template.parentheses(gs_table.iloc[loc, 31]))[0][0]
            cicc_rating = issuer.iloc[location, 22]
            if gs_table.iloc[loc, 8] is not None or not math.isnan(gs_table.iloc[loc, 8]):
                row[1].text = str(cicc_rating)
            else:
                row[1].text = "可能被列入名单"
                print("请在issue表手动查阅中金评级，可能被列入名单")

    def get_industry(gs_table,issuer,issuer_column,row,loc):
        if len(np.where(issuer_column == doc_template.parentheses(gs_table.iloc[loc, 31]))[0]) == 0:
            row[1].text = "手动查找"
            print("%s的行业在issue表未找到" % (gs_table.iloc[loc, 1]))
        else:
            location = np.where(issuer_column == doc_template.parentheses(gs_table.iloc[loc, 31]))[0][0]
            industry = issuer.iloc[location, 20]
            row[1].text = industry

    def parentheses(string):    #用于在CP issue index里找主体名称（因为wind导出的主体全称有括号的是英文括号而CP issue index里都是中文括号）
        English = u'()'
        Chinese = u'（）'
        table = {ord(o): ord(k) for o, k in zip(English, Chinese)}
        return string.translate(table)

    def table_modify(self,type=None,auto_detect=True):  #支持手动输入type，但需要auto_detect=False
        doc = Document(self.targetpath)
        gs_table = pd.read_excel(self.excelpath)
        fill="XXX"
        gs_table.iloc[:,31],gs_table.iloc[:,28],gs_table.iloc[:,53],gs_table.iloc[:,41],gs_table.iloc[:,10]=gs_table.iloc[:,31].fillna(fill),gs_table.iloc[:,28].fillna(fill),gs_table.iloc[:,53].fillna(fill),gs_table.iloc[:,41].fillna(fill),gs_table.iloc[:,10].fillna(fill)
        issuer = pd.read_excel(self.issuerpath)  # 导入比较慢
        issuer_column = issuer.iloc[:, 2]
        if auto_detect:
            if str(gs_table.iloc[1,65]).find("中")!=-1 or str(gs_table.iloc[1,65]).find("短")!=-1:
                type="CPMTN"
            else:
                type="CB"
        for ind in range(len(doc.tables)):
            if ind == 0 or ind == len(doc.tables) - 1 or ind == len(doc.tables) - 2:  # 去掉头尾table
                continue
            table = doc.tables[ind]
            bondnamelist = []
            for index in range(len(table.rows)):
                row = table.rows[index].cells     #这是某个表里的第几行
                if index == 0 or index == len(table.rows) - 1:  # 去掉头尾空行
                    continue
                if index == 1:
                    doc_template.get_bond_name(bondnamelist,ind,row,gs_table)
                if index == 2:  # 发行人
                    row[1].text = gs_table.iloc[loc, 31]
                if index == 3:  # 债券期限
                    doc_template.get_length(loc,row,gs_table)
                if index == 4:  # 发行额
                    if math.isnan(gs_table.iloc[loc, 5]):
                        row[1].text = str(gs_table.iloc[loc, 4])+"亿元人民币"  # 目前发行额只能有一个空
                    else:
                        row[1].text = str(gs_table.iloc[loc, 5])+"亿元人民币"
                if index == 5 and type=="CB":     #发行利率
                    row[1].text = "待确定"
                if index ==5 and type=="CPMTN":    #付息方式
                    if gs_table.iloc[loc, 65].find("中")<0:
                        row[1].text = "到期还本付息"
                    else:
                        row[1].text = "周期性付息"
                if index == 6 and type=="CB":  # 增信方式
                    doc_template.get_credit_boost(self,row,loc,gs_table)
                if index == 6 and type=="CPMTN":
                    row[1].text = "待确定"
                if index == 7 and type=="CB":  # 信用级别
                    doc_template.get_ratings(self,row,gs_table,type)
                if index == 7 and type=="CPMTN":
                    doc_template.get_credit_boost(self,row, loc, gs_table)
                if index == 8 and type=="CB":  # 发行场所
                    if gs_table.iloc[loc, 53][0]=="上":
                       row[1].text = "上交所"
                    elif gs_table.iloc[loc, 53][0]=="深":
                       row[1].text = "深交所"
                    elif gs_table.iloc[loc, 53][0]=="银":
                       row[1].text = "银行间"
                    else:
                        row[1].text = gs_table.iloc[loc, 53]
                if index ==8 and type=="CPMTN":  # 信用级别
                    doc_template.get_ratings(self, row, gs_table, type)
                if index == 9:  # 主承销商
                    row[1].text = str(gs_table.iloc[loc, 41])
                if index == 10:  # 发行日
                    row[1].text = gs_table.iloc[loc, 2].strftime('%Y{y}%m{m}%d{d}').format(y='年', m='月', d='日')
                if index == 11 and type=="CB":  # 行业
                    doc_template.get_industry(gs_table, issuer,issuer_column, row, loc)
                if index == 11 and type=="CPMTN":
                    row[1].text = gs_table.iloc[loc, 20].strftime('%Y{y}%m{m}%d{d}').format(y='年', m='月', d='日')
                if index == 12 and type == "CB":  # 中金评级
                    doc_template.get_CICC_ratings(row,loc,gs_table,issuer_column,issuer)
                if index == 12 and type == "CPMTN":
                    doc_template.get_industry(gs_table,issuer,issuer_column,row,loc)
                if index == 13 and type == "CPMTN":   #中金评级
                    doc_template.get_CICC_ratings(row, loc, gs_table, issuer_column, issuer)
                else: continue
            else: continue
        doc.save(self.targetpath)
        print("Done")
